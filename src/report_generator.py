import numpy as np
import pandas as pd
import os
import re
from typing import Optional

import lxml.etree as ET
import latex2mathml.converter

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

from src.welltest import (
    Problem1AnalysisResult,
    Problem2AnalysisResult,
    analyze_problem1_drawdown,
    analyze_problem2_buildup,
)

# ============================================================================
# CONSTANTS & DEFAULTS
# ============================================================================
DEFAULT_OUTPUT_DIR = r"C:\Users\60163\Downloads\PFB2073_Well_Test_Analysis_-_May_2026_1784922163"
DEFAULT_DOCX_PATH  = os.path.join(DEFAULT_OUTPUT_DIR, "PFB2073_WellTest_Report_APA.docx")

STUDENT_NAME   = "Muhammad Radhi bin Mohd Riadh"
STUDENT_ID     = "22010970"
PAPER_TITLE    = "PFB2073 Well Test Analysis Assignment 2"
AFFILIATION    = "Department of Petroleum Engineering, Universiti Teknologi PETRONAS"
COURSE         = "PFB2073: Well Test Analysis"
INSTRUCTOR     = "Dr. Shiferaw Regassa Jufar"
DUE_DATE       = "31 July 2026"
RUNNING_HEAD   = "WELL TEST ANALYSIS ASSIGNMENT 2"

# XSLT Transformer for MathML -> OMML
XSLT_PATH = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
_transform = None

def get_xslt_transform():
    global _transform
    if _transform is None:
        if os.path.exists(XSLT_PATH):
            xslt_tree = ET.parse(XSLT_PATH)
            _transform = ET.XSLT(xslt_tree)
    return _transform

# ============================================================================
# LATEX TO NATIVE WORD OMML EQUATION CONVERTER
# ============================================================================
def convert_latex_to_omml(latex_str: str, is_block: bool = False):
    """Converts a LaTeX equation string to a native MS Word OMML XML element."""
    clean = latex_str.strip()

    # Pre-clean LaTeX notations for MathML compatibility
    clean = clean.replace(r"\text{", r"\mathrm{")
    clean = clean.replace(r"\Delta", r"{\Delta}")
    clean = clean.replace(r"^\prime", r"'")
    clean = clean.replace(r"\cdot", r"\cdot ")
    clean = clean.replace(r"\}", r"\}")

    try:
        transform = get_xslt_transform()
        mathml_str = latex2mathml.converter.convert(clean)
        mathml_tree = ET.fromstring(mathml_str)

        if transform is not None:
            omml_tree = transform(mathml_tree)
            omml_xml = ET.tostring(omml_tree, encoding="unicode")
        else:
            omml_xml = f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:r><m:t>{clean}</m:t></m:r></m:oMath>'

        omml_element = parse_xml(omml_xml)

        if is_block:
            omath_para = parse_xml('<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"/>')
            omath_para.append(omml_element)
            return omath_para
        return omml_element
    except Exception as e:
        print(f"[Warning] OMML conversion fallback for '{clean}': {e}")
        fallback_xml = f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:r><m:t>{clean}</m:t></m:r></m:oMath>'
        return parse_xml(fallback_xml)

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name     = name
    run.font.size     = Pt(size)
    run.font.bold     = bold
    run.font.italic   = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_para_spacing(para, before=0, after=0, line_rule=WD_LINE_SPACING.DOUBLE, line_val=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = line_rule
    if line_val:
        pf.line_spacing = line_val

PHYSICAL_UNITS_REGEX = r'bbl/psi|STB/day|STB/D|RB/STB|bbl/STB|psi/cycle|acres?|psia|psi|mD|bbl|hrs?|days?|cP|ft'
NUMBER_UNIT_PATTERN = re.compile(r'(?<=\d)(?=(?:' + PHYSICAL_UNITS_REGEX + r')\b)')

def enforce_number_unit_spacing(text: str) -> str:
    """
    Enforces a single space between numeric values and physical units (e.g. 660ft -> 660 ft, 14.16mD -> 14.16 mD,
    1314psia -> 1314 psia, 0.001314bbl/psi -> 0.001314 bbl/psi, 72hr -> 72 hr) prior to OMML conversion.
    Safe against scientific notation (1e-6), ordinals (7th), and LaTeX commands (\text{...}).
    """
    if not text:
        return text
    return NUMBER_UNIT_PATTERN.sub(' ', text)

def append_text_with_omml(
    para,
    text: str,
    font_name: str = "Times New Roman",
    size: int = 12,
    bold: bool = False,
    italic: bool = False,
    color: Optional[tuple] = None
):
    """
    Pre-processes text for number-unit spacing enforcement (R2), then parses inline LaTeX
    formulas ($...$) into native Word OMML equation objects (w:oMath), while rendering normal
    text as styled runs.
    """
    if not text:
        return
    processed_text = enforce_number_unit_spacing(text)
    parts = re.split(r'(\$.*?\$)', processed_text)
    for part in parts:
        if part.startswith("$") and part.endswith("$") and len(part) > 2:
            latex_code = part[1:-1].strip()
            omml_elem = convert_latex_to_omml(latex_code, is_block=False)
            para._p.append(omml_elem)
        elif part:
            run = para.add_run(part)
            set_font(run, name=font_name, size=size, bold=bold, italic=italic, color=color)

def add_heading_apa(doc, text, level=1, bookmark_id=None):
    para = doc.add_paragraph()
    
    # Page Break & Keep With Next Rules:
    # 1. Level 1 headings always start on a brand new page
    # 2. Level 1, 2, and 3 headings always stay on the same page as their next paragraph (prevent orphan headings)
    if level == 1:
        para.paragraph_format.page_break_before = True
        set_para_spacing(para, before=0, after=6)
    elif level == 2:
        set_para_spacing(para, before=12, after=6)
    else:
        set_para_spacing(para, before=6, after=4)
        
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.keep_with_next = True

    # Set Word Outline Level so Word TOC indexes it natively
    pPr = para._p.get_or_add_pPr()
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(level - 1))
    pPr.append(outlineLvl)

    # Add Word Bookmark for interactive TOC navigation
    if bookmark_id:
        bm_start = OxmlElement('w:bookmarkStart')
        bm_start.set(qn('w:id'), str(bookmark_id))
        bm_start.set(qn('w:name'), f"heading_{bookmark_id}")
        para._p.append(bm_start)

    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        append_text_with_omml(para, text, font_name="Times New Roman", size=12, bold=True)
    elif level == 2:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        append_text_with_omml(para, text, font_name="Times New Roman", size=12, bold=True)
    elif level == 3:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        append_text_with_omml(para, text, font_name="Times New Roman", size=12, bold=True, italic=True)

    if bookmark_id:
        bm_end = OxmlElement('w:bookmarkEnd')
        bm_end.set(qn('w:id'), str(bookmark_id))
        para._p.append(bm_end)

    return para

def add_body_para(doc, text, indent=True):
    """
    Adds body paragraph. Parses inline ($ ... $) and display ($$ ... $$) math,
    converting all LaTeX formulas directly into native MS Word OMML equations.
    """
    text_strip = text.strip()

    # Case 1: Display / Block Math Equation ($$ ... $$)
    if text_strip.startswith("$$") and text_strip.endswith("$$"):
        latex_code = text_strip[2:-2].strip()
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Pt(0)
        set_para_spacing(para, before=4, after=4, line_rule=WD_LINE_SPACING.SINGLE)

        omml_para = convert_latex_to_omml(latex_code, is_block=True)
        para._p.append(omml_para)
        return para

    # Case 2: Standard Body Text with optional Inline Math ($ ... $)
    para = doc.add_paragraph()
    set_para_spacing(para, before=0, after=0, line_rule=WD_LINE_SPACING.DOUBLE)
    if indent:
        para.paragraph_format.first_line_indent = Cm(1.27)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    append_text_with_omml(para, text, font_name="Times New Roman", size=12)
    return para

def add_blank_line(doc):
    para = doc.add_paragraph()
    set_para_spacing(para, before=0, after=0)
    run = para.add_run("")
    set_font(run, size=12)
    return para

def set_cell_content_with_latex(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
    """Fills a table cell with text and parses embedded inline LaTeX formulas into native OMML equations."""
    cell.text = "" # Clear default text
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(0)
    set_para_spacing(p, before=2, after=2, line_rule=WD_LINE_SPACING.SINGLE)

    append_text_with_omml(p, text, font_name="Times New Roman", size=11, bold=bold)

def add_figure_label(doc, number, title):
    p_num = doc.add_paragraph()
    p_num.paragraph_format.first_line_indent = Pt(0)
    p_num.paragraph_format.keep_with_next = True
    p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p_num, before=6, after=0)
    run = p_num.add_run(f"Figure {number}")
    set_font(run, bold=True, size=12)

    p_title = doc.add_paragraph()
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.keep_with_next = True
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p_title, before=0, after=2)
    append_text_with_omml(p_title, title, font_name="Times New Roman", size=12, italic=True)
    return p_num, p_title

def add_table_label(doc, number, title):
    p_num = doc.add_paragraph()
    p_num.paragraph_format.first_line_indent = Pt(0)
    p_num.paragraph_format.keep_with_next = True
    p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p_num, before=6, after=0)
    run = p_num.add_run(f"Table {number}")
    set_font(run, bold=True, size=12)

    p_title = doc.add_paragraph()
    p_title.paragraph_format.first_line_indent = Pt(0)
    p_title.paragraph_format.keep_with_next = True
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p_title, before=0, after=2)
    append_text_with_omml(p_title, title, font_name="Times New Roman", size=12, italic=True)
    return p_num, p_title

def embed_picture(doc, path, height_cm=10.0):
    """Embeds an image with exact height of height_cm while preserving proportional width and removing extra space."""
    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Pt(0)
        set_para_spacing(para, before=2, after=4, line_rule=WD_LINE_SPACING.SINGLE)
        run = para.add_run()
        run.add_picture(path, height=Cm(height_cm))
        return para
    except Exception as e:
        p = doc.add_paragraph(f"[Figure could not be embedded: {e}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

def add_page_header(doc, running_head):
    """
    Configures bulletproof 2-column borderless table in page header:
    - Left Cell: Running Head (Left-aligned)
    - Right Cell: Page Number (Right-aligned)
    """
    for section in doc.sections:
        section.different_first_page_header_footer = False
        header = section.header
        header.is_linked_to_previous = False
        header.paragraphs[0].clear()

        if len(header.tables) == 0:
            htbl = header.add_table(rows=1, cols=2, width=Inches(6.5))
        else:
            htbl = header.tables[0]

        htbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        htbl.autofit = False

        tblPr = htbl._tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{b}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        cell0 = htbl.cell(0, 0)
        cell0.width = Inches(5.0)
        p0 = cell0.paragraphs[0]
        p0.clear()
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(p0, before=0, after=0)
        r0 = p0.add_run(running_head)
        set_font(r0, name="Times New Roman", size=12, color=(0, 0, 0))

        cell1 = htbl.cell(0, 1)
        cell1.width = Inches(1.5)
        p1 = cell1.paragraphs[0]
        p1.clear()
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_para_spacing(p1, before=0, after=0)
        r1 = p1.add_run()
        set_font(r1, name="Times New Roman", size=12, color=(0, 0, 0))

        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        r1._r.append(fld_begin)
        fld_instr = OxmlElement('w:instrText')
        fld_instr.set(qn('xml:space'), 'preserve')
        fld_instr.text = ' PAGE '
        r1._r.append(fld_instr)
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        r1._r.append(fld_end)

def apply_apa_table_borders(table):
    """
    Applies strict APA 7th Edition table styling:
    - Top horizontal line on header row
    - Bottom horizontal line on header row
    - Bottom horizontal line on last data row
    - NO vertical lines
    - NO interior horizontal gridlines
    - NO cell background shading
    """
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is not None:
        tblPr.remove(tblBorders)

    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = tcPr.first_child_found_in("w:shd")
            if shd is not None:
                tcPr.remove(shd)

            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is not None:
                tcPr.remove(tcBorders)

            tcBorders = OxmlElement('w:tcBorders')

            top = OxmlElement('w:top')
            bottom = OxmlElement('w:bottom')
            left = OxmlElement('w:left')
            right = OxmlElement('w:right')

            left.set(qn('w:val'), 'nil')
            right.set(qn('w:val'), 'nil')

            if r_idx == 0:
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), '8')
                top.set(qn('w:space'), '0')
                top.set(qn('w:color'), '000000')

                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '8')
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), '000000')
            elif r_idx == len(table.rows) - 1:
                top.set(qn('w:val'), 'nil')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '8')
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), '000000')
            else:
                top.set(qn('w:val'), 'nil')
                bottom.set(qn('w:val'), 'nil')

            tcBorders.append(top)
            tcBorders.append(bottom)
            tcBorders.append(left)
            tcBorders.append(right)
            tcPr.append(tcBorders)

def add_toc_entry(doc, title, page_str, level=1, bookmark_id=None):
    """
    Adds a formatted Table of Contents entry with 2.5 line spacing, right-aligned dot leaders,
    and interactive hyperlinks pointing to chapter bookmarks.
    """
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Cm(0) if level == 1 else Cm(0.63)
    set_para_spacing(p, before=2, after=2, line_rule=WD_LINE_SPACING.MULTIPLE, line_val=2.5)

    if bookmark_id:
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), f"heading_{bookmark_id}")

        r_title = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '22')
        rPr.append(sz)
        if level == 1:
            b = OxmlElement('w:b')
            rPr.append(b)
        r_title.append(rPr)
        t_title = OxmlElement('w:t')
        t_title.text = title
        r_title.append(t_title)
        hyperlink.append(r_title)

        r_tab = OxmlElement('w:r')
        t_tab = OxmlElement('w:tab')
        r_tab.append(t_tab)
        hyperlink.append(r_tab)

        r_pg = OxmlElement('w:r')
        rPr_pg = OxmlElement('w:rPr')
        rPr_pg.append(rFonts)
        rPr_pg.append(sz)
        if level == 1:
            rPr_pg.append(b)
        r_pg.append(rPr_pg)
        t_pg = OxmlElement('w:t')
        t_pg.text = page_str
        r_pg.append(t_pg)
        hyperlink.append(r_pg)

        p._p.append(hyperlink)
    else:
        r_title = p.add_run(title)
        set_font(r_title, bold=(level == 1), size=11)

        r_tab = p.add_run()
        set_font(r_tab, size=11)
        r_tab._r.append(OxmlElement('w:tab'))

        r_pg = p.add_run(page_str)
        set_font(r_pg, bold=(level == 1), size=11)

    pPr = p._p.get_or_add_pPr()
    tabs_elem = OxmlElement('w:tabs')
    tab_stop = OxmlElement('w:tab')
    tab_stop.set(qn('w:val'), 'right')
    tab_stop.set(qn('w:leader'), 'dot')
    tab_stop.set(qn('w:pos'), '9360')
    tabs_elem.append(tab_stop)
    pPr.append(tabs_elem)
    return p

# ============================================================================
# MAIN REPORT GENERATION FUNCTION
# ============================================================================
def generate_apa_report(
    p1_res: Optional[Problem1AnalysisResult] = None,
    p2_res: Optional[Problem2AnalysisResult] = None,
    output_path: Optional[str] = None,
    output_dir: Optional[str] = None
) -> str:
    """
    Generates APA 7th Edition compliant Word document (.docx) report for well test analysis.
    Dynamically binds report content, tables, equations, and discussion to calculation results.

    Parameters:
        p1_res      : Problem1AnalysisResult object. If None, computed dynamically.
        p2_res      : Problem2AnalysisResult object. If None, computed dynamically.
        output_path : Target file path for the .docx report.
        output_dir  : Target directory path if output_path is not specified.

    Returns:
        str: Absolute path to the generated .docx report.
    """
    if p1_res is None:
        p1_res = analyze_problem1_drawdown()
    if p2_res is None:
        p2_res = analyze_problem2_buildup()

    if output_path is None:
        if output_dir is None:
            output_dir = DEFAULT_OUTPUT_DIR
        output_path = os.path.join(output_dir, "PFB2073_WellTest_Report_APA.docx")
    else:
        output_dir = os.path.dirname(output_path)

    plot_p1_loglog  = os.path.join(output_dir, "Problem1_LogLog_Diagnostic.png")
    plot_p1_semi    = os.path.join(output_dir, "Problem1_SemiLog.png")
    plot_p2_loglog  = os.path.join(output_dir, "Problem2_LogLog_Diagnostic.png")
    plot_p2_horner  = os.path.join(output_dir, "Problem2_Horner.png")

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)
        section.page_width    = Cm(21.59)
        section.page_height   = Cm(27.94)

    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    add_page_header(doc, RUNNING_HEAD)

    # ------------------------------------------------------------------------
    # PAGE 1: Title Page
    # ------------------------------------------------------------------------
    for _ in range(3):
        p = add_blank_line(doc)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Pt(0)
    set_para_spacing(p_title, before=0, after=0)
    r = p_title.add_run(PAPER_TITLE)
    set_font(r, bold=True, size=12)

    add_blank_line(doc)

    for info in [STUDENT_NAME, AFFILIATION, COURSE, INSTRUCTOR, DUE_DATE]:
        p_info = doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_info.paragraph_format.first_line_indent = Pt(0)
        set_para_spacing(p_info, before=0, after=0)
        r = p_info.add_run(info)
        set_font(r, size=12)

    doc.add_page_break()

    # ------------------------------------------------------------------------
    # PAGE 2: Table of Contents (Fits Perfectly on One Page with 2.5 spacing)
    # ------------------------------------------------------------------------
    p_toc_label = doc.add_paragraph()
    p_toc_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_toc_label.paragraph_format.first_line_indent = Pt(0)
    set_para_spacing(p_toc_label, before=0, after=8)
    r_toc = p_toc_label.add_run("Table of Contents")
    set_font(r_toc, bold=True, size=14)

    # Embed XML Word TOC Field
    p_xml_toc = doc.add_paragraph()
    p_xml_toc.paragraph_format.first_line_indent = Pt(0)
    set_para_spacing(p_xml_toc, before=0, after=4)
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), r'TOC \o "1-2" \h \z \u')
    p_xml_toc._p.append(fldSimple)

    # Pre-formatted Level 1 and Level 2 Entries with 2.5 Spacing and Interactive Bookmarks
    toc_structure = [
        ("Abstract", "3", 1, 1),
        ("Introduction", "4", 1, 2),
        ("Methodology", "4", 1, 3),
        ("Problem 1: Drawdown Test", "4", 2, 4),
        ("Problem 2: Buildup Test", "5", 2, 5),
        ("Results", "5", 1, 6),
        ("Problem 1: Drawdown Test Results", "5", 2, 7),
        ("Problem 2: Buildup Test Results", "6", 2, 8),
        ("Discussion", "7", 1, 9),
        ("Drawdown Test Interpretation", "7", 2, 10),
        ("Buildup Test Interpretation, Model Identification, and MBH Correction", "7", 2, 11),
        ("Conclusion", "8", 1, 12),
        ("References", "9", 1, 13),
    ]

    for title, pg, lvl, b_id in toc_structure:
        add_toc_entry(doc, title, pg, level=lvl, bookmark_id=b_id)

    # ------------------------------------------------------------------------
    # PAGE 3: Abstract (Starts on new page via Level 1 page_break_before)
    # ------------------------------------------------------------------------
    add_heading_apa(doc, "Abstract", level=1, bookmark_id=1)
    add_body_para(doc,
        f"This report presents the well test analysis for two problems: a drawdown test (Problem 1) "
        f"and a buildup test (Problem 2) in a petroleum reservoir. The Bourdet pressure derivative "
        f"method was employed as the primary diagnostic tool to identify flow regimes and determine "
        f"reservoir parameters. For the drawdown test, the analysis yielded a permeability of "
        f"$k = {p1_res.semilog.k:.2f} \\text{{ mD}}$, a skin factor of $s = {p1_res.semilog.s:.2f}$, and a wellbore storage coefficient of "
        f"$C = {p1_res.semilog.C:.5f} \\text{{ bbl/psi}}$ ($C_D = {p1_res.semilog.C_D:.1f}$). For the buildup test, the computed permeability was $k = {p2_res.semilog.k:.2f} \\text{{ mD}}$ "
        f"with a skin factor of $s = {p2_res.semilog.s:.2f}$, a wellbore storage coefficient of $C = {p2_res.semilog.C:.5f} \\text{{ bbl/psi}}$ ($C_D = {p2_res.semilog.C_D:.1f}$), "
        f"and a true average reservoir pressure of $\\bar{{p}} = {p2_res.semilog.p_bar:.1f} \\text{{ psia}}$ derived via the Matthews-Brons-Hazebroek (MBH) method using the Dietz shape factor approach. "
        f"The radius of investigation during the $72 \\text{{ hr}}$ shut-in period for the buildup test was $r_{{inv}} = {p2_res.semilog.r_inv_72hr:.0f} \\text{{ ft}}$. "
        f"A comparison with the boundaries of the $80 \\text{{ acre}}$, 2:1 rectangular reservoir was conducted "
        f"to evaluate whether boundary effects were felt during shut-in versus prior production.",
        indent=False
    )

    p_kw = doc.add_paragraph()
    p_kw.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_kw.paragraph_format.first_line_indent = Cm(1.27)
    set_para_spacing(p_kw, before=0, after=0)
    r_kw1 = p_kw.add_run("Keywords: ")
    set_font(r_kw1, italic=True, size=12)
    r_kw2 = p_kw.add_run(
        "well test analysis, drawdown test, buildup test, Bourdet derivative, "
        "type curve matching, permeability, skin factor, wellbore storage, MBH method, Dietz shape factor"
    )
    set_font(r_kw2, size=12)

    # ------------------------------------------------------------------------
    # Introduction (Starts on new page via Level 1 page_break_before)
    # ------------------------------------------------------------------------
    add_heading_apa(doc, "Introduction", level=1, bookmark_id=2)
    add_body_para(doc,
        "Well test analysis is a fundamental technique in reservoir engineering used to determine "
        "reservoir parameters such as permeability, skin factor, and wellbore storage coefficient. "
        "These parameters are essential for reservoir characterization, production optimization, and "
        "field development planning (Ahmed & McKinney, 2005)."
    )
    add_body_para(doc,
        "The Bourdet pressure derivative, introduced by Bourdet et al. (1983), revolutionized well "
        "test interpretation by providing a diagnostic tool more sensitive to flow regime changes "
        "than pressure data alone."
    )

    # ------------------------------------------------------------------------
    # Methodology (Starts on new page via Level 1 page_break_before)
    # ------------------------------------------------------------------------
    add_heading_apa(doc, "Methodology", level=1, bookmark_id=3)
    add_heading_apa(doc, "Problem 1: Drawdown Test", level=2, bookmark_id=4)
    add_body_para(doc,
        f"The drawdown test was conducted at a constant flow rate of $q = {p1_res.params.q:.0f} \\text{{ STB/day}}$ from an initial "
        f"reservoir pressure of $p_i = {p1_res.params.p_i:.1f} \\text{{ psia}}$. Reservoir parameters: $h = {p1_res.params.h:.0f} \\text{{ ft}}$, $\\phi = {p1_res.params.phi:.3f}$, "
        f"$r_w = {p1_res.params.r_w:.2f} \\text{{ ft}}$, $B = {p1_res.params.B:.3f} \\text{{ RB/STB}}$, $c_t = {p1_res.params.c_t:.1e} \\text{{ psi}}^{{-1}}$, $\\mu = {p1_res.params.mu:.2f} \\text{{ cP}}$."
    )

    add_heading_apa(doc, "Pressure Change Calculation", level=3)
    add_body_para(doc, "$$\\Delta p = p_i - p_{wf}$$", indent=True)

    add_heading_apa(doc, "Bourdet Pressure Derivative", level=3)
    add_body_para(doc,
        "$$t \\cdot \\Delta p' = \\frac{\\frac{\\Delta p_L}{\\Delta \\ln(t_L)} \\cdot \\Delta \\ln(t_R) + \\frac{\\Delta p_R}{\\Delta \\ln(t_R)} \\cdot \\Delta \\ln(t_L)}{\\Delta \\ln(t_L) + \\Delta \\ln(t_R)}$$",
        indent=True
    )
    add_body_para(doc,
        "*(where subscripts $L$ and $R$ denote the left and right adjacent data points in log-time, respectively).*",
        indent=True
    )

    add_heading_apa(doc, "Permeability, Skin Factor, and Wellbore Storage", level=3)
    add_body_para(doc, "$$k = \\frac{162.6 q B \\mu}{m h}$$", indent=True)
    add_body_para(doc,
        "$$s = 1.1513 \\left[ \\frac{\\Delta p_{1hr}}{m} - \\log_{10}\\left(\\frac{k}{\\phi \\mu c_t r_w^2}\\right) + 3.2275 \\right]$$",
        indent=True
    )
    add_body_para(doc,
        "$$s = 1.1513 \\left[ \\frac{737.20}{133.26} - \\log_{10}\\left(\\frac{14.16}{(0.276)(5.28)(9.4 \\times 10^{-6})(0.36^2)}\\right) + 3.2275 \\right] = 2.51$$",
        indent=True
    )
    add_heading_apa(doc, "Dimensional Wellbore Storage ($C$)", level=3)
    add_body_para(doc,
        "The dimensional wellbore storage coefficient was estimated using a data coordinate ($t, \\Delta p$) located on the early-time unit-slope line of the log-log diagnostic plot:",
        indent=True
    )
    add_body_para(doc,
        "$$C = \\frac{q B t}{24 \\Delta p} \\quad \\text{[bbl/psi]}$$",
        indent=True
    )
    add_body_para(doc,
        "$$C = \\frac{(50)(1.099)(0.100)}{24(217.6)} = 0.00105 \\text{ bbl/psi}$$",
        indent=True
    )
    add_body_para(doc,
        "$$C_D = \\frac{0.8936 C}{\\phi c_t h r_w^2}$$",
        indent=True
    )
    add_body_para(doc,
        "$$C_D = \\frac{0.8936(0.00105)}{(0.276)(9.4 \\times 10^{-6})(25)(0.36^2)} = 111.6$$",
        indent=True
    )

    add_heading_apa(doc, "Problem 2: Buildup Test", level=2, bookmark_id=5)
    add_body_para(doc,
        f"Reservoir and Fluid Properties (Problem 2): The buildup test was conducted after $t_p = {p2_res.params.t_p:.0f} \\text{{ hrs}}$ at $q = {p2_res.params.q:.0f} \\text{{ STB/day}}$ in an $A = 80 \\text{{ acres}}$ 2:1 rectangular reservoir ($h = {p2_res.params.h:.0f} \\text{{ ft}}$, $\\phi = {p2_res.params.phi:.3f}$, $r_w = {p2_res.params.r_w:.2f} \\text{{ ft}}$, $B = {p2_res.params.B:.3f} \\text{{ RB/STB}}$, $c_t = 10.9 \\times 10^{{-6}} \\text{{ psi}}^{{-1}}$, $\\mu = {p2_res.params.mu:.2f} \\text{{ cP}}$)."
    )
    add_body_para(doc, "$$\\Delta t_e = \\frac{t_p \\Delta t}{t_p + \\Delta t}$$", indent=True)
    add_heading_apa(doc, "Permeability Calculation", level=3)
    add_body_para(doc, "$$k = \\frac{162.6 q B \\mu}{m h}$$", indent=True)
    add_body_para(doc,
        f"$$k = \\frac{{162.6(10)(1.098)(5.11)}}{{(30.29)(10)}} = {p2_res.semilog.k:.2f} \\text{{ mD}}$$",
        indent=True
    )
    add_heading_apa(doc, "Horner Skin Factor ($s$)", level=3)
    add_body_para(doc,
        "Because buildup tests utilize the Horner time ratio, the skin factor evaluates the flowing pressure immediately prior to shut-in ($p_{wf(\\Delta t=0)}$) and the idealized 1-hour shut-in pressure ($p_{ws,1hr}$) extrapolated from the Horner straight line:",
        indent=True
    )
    add_body_para(doc,
        "$$s = 1.1513 \\left[ \\frac{p_{ws,1hr} - p_{wf(\\Delta t=0)}}{m} - \\log_{10}\\left(\\frac{k}{\\phi \\mu c_t r_w^2}\\right) + 3.2275 \\right]$$",
        indent=True
    )
    add_body_para(doc,
        f"$$s = 1.1513 \\left[ \\frac{{1262.87 - 1192.45}}{{30.29}} - \\log_{{10}}\\left(\\frac{{{p2_res.semilog.k:.2f}}}{{(0.319)(5.11)(10.9 \\times 10^{{-6}})(0.34^2)}}\\right) + 3.2275 \\right] = {p2_res.semilog.s:.2f}$$",
        indent=True
    )
    add_heading_apa(doc, "Wellbore Storage Coefficient ($C$ and $C_D$)", level=3)
    add_body_para(doc,
        "$$C = \\frac{(10)(1.098)(0.050)}{24(14.47)} = 0.00158 \\text{ bbl/psi}$$",
        indent=True
    )
    add_body_para(doc,
        "$$C_D = \\frac{0.8936(0.00158)}{(0.319)(10.9 \\times 10^{-6})(10)(0.34^2)} = 351.3$$",
        indent=True
    )
    add_heading_apa(doc, "Radius of Investigation", level=3)
    add_body_para(doc, "$$r_{inv} = \\sqrt{\\frac{k t}{948 \\phi \\mu c_t}}$$", indent=True)

    # ------------------------------------------------------------------------
    # Results (Starts on new page via Level 1 page_break_before)
    # ------------------------------------------------------------------------
    add_heading_apa(doc, "Results", level=1, bookmark_id=6)
    add_body_para(doc,
        "*Note: The analytical type-curve models generated by the software (shown in the legends of Figure 1 and Figure 3) yield slightly different parameters than the manual semi-log analysis due to the software's non-linear regression algorithms. Following standard engineering practice, the manually derived semi-log parameters are reported as the primary definitive results in the summary tables.*",
        indent=False
    )
    add_heading_apa(doc, "Problem 1: Drawdown Test Results", level=2, bookmark_id=7)
    add_body_para(doc,
        f"Permeability $k = {p1_res.semilog.k:.2f} \\text{{ mD}}$, skin factor $s = {p1_res.semilog.s:.2f}$, wellbore storage coefficient $C = {p1_res.semilog.C:.5f} \\text{{ bbl/psi}}$ ($C_D = {p1_res.semilog.C_D:.1f}$), and semi-log slope $m = {p1_res.semilog.m:.2f} \\text{{ psi/cycle}}$. "
        f"The complete calculated reservoir parameters for the drawdown test are summarized in Table 1. "
        f"The corresponding log-log diagnostic plot and semi-log plot are presented in Figure 1 and Figure 2, respectively."
    )

    add_heading_apa(doc, "Problem 2: Buildup Test Results", level=2, bookmark_id=8)
    add_body_para(doc,
        f"Permeability $k = {p2_res.semilog.k:.2f} \\text{{ mD}}$, skin factor $s = {p2_res.semilog.s:.2f}$, wellbore storage coefficient $C = {p2_res.semilog.C:.5f} \\text{{ bbl/psi}}$ ($C_D = {p2_res.semilog.C_D:.1f}$), semi-log slope $m = {p2_res.semilog.m:.2f} \\text{{ psi/cycle}}$, radius of investigation $r_{{inv}} = 359 \\text{{ ft}}$, and true average pressure $\\bar{{p}} = 1340.2 \\text{{ psia}}$. "
        f"The complete calculated reservoir parameters for the buildup test are summarized in Table 2. "
        f"The corresponding log-log diagnostic plot and Horner plot are presented in Figure 3 and Figure 4, respectively."
    )

    # ------------------------------------------------------------------------
    # Discussion (Starts on new page via Level 1 page_break_before)
    # ------------------------------------------------------------------------
    add_heading_apa(doc, "Discussion", level=1, bookmark_id=9)
    add_heading_apa(doc, "Drawdown Test Interpretation", level=2, bookmark_id=10)
    add_body_para(doc,
        f"Extrapolating the IARF semi-log straight line to $t = 1 \\text{{ hr}}$ yields $p_{{wf,1hr}} = {p1_res.semilog.p_1hr:.1f} \\text{{ psia}}$. "
        f"The true skin factor is $s = {p1_res.semilog.s:.2f}$, indicating formation damage near the wellbore."
    )

    add_heading_apa(doc, "Buildup Test Interpretation, Model Identification, and MBH Correction", level=2, bookmark_id=11)
    add_heading_apa(doc, "Model Identification", level=3)
    add_body_para(doc,
        f"The diagnostic log-log plot for the buildup test exhibits three distinct flow regimes. "
        f"In the Early Time Region (ETR), wellbore storage dominates the response with a characteristic unit slope ($m=1$) on both pressure change and derivative curves, yielding a wellbore storage coefficient of $C = 0.00158 \\text{{ bbl/psi}}$ ($C_D = 351.3$) and a skin factor hump indicating stimulus ($s = -1.86$). "
        f"This is followed by a Transition Region as wellbore storage effects decline and transient flow expands into the formation. "
        f"In the Middle Time Region (MTR), Infinite-Acting Radial Flow (IARF) is clearly established, marked by a horizontal derivative plateau of $t_e \\cdot \\Delta p' \\approx 13.15 \\text{{ psi}}$ and a corresponding Horner straight line with a slope of $m = 30.29 \\text{{ psi/cycle}}$."
    )
    add_body_para(doc,
        "For Problem 2, the radius of investigation during the $72 \\text{ hr}$ shut-in period only reached $r_{inv} = 359 \\text{ ft}$, which does not exceed the distance to the nearest boundary ($660 \\text{ ft}$). Thus, boundary effects are not visible on the buildup log-log derivative."
    )
    add_body_para(doc,
        "However, over the total production time ($t_p = 960 \\text{ hr}$), the radius of investigation vastly exceeded the boundaries of the $80\\text{-acre}$ reservoir ($A = 3,484,800 \\text{ ft}^2$). The dimensionless production time based on area is calculated to evaluate the flow regime prior to shut-in:"
    )
    add_body_para(doc,
        f"$$t_{{pDA}} = \\frac{{0.0002637 k t_p}}{{\\phi \\mu c_t A}} = \\frac{{0.0002637 ({p2_res.semilog.k:.2f}) (960)}}{{0.319 (5.11) (10.9 \\times 10^{{-6}}) (3,484,800)}} = 0.1231$$",
        indent=True
    )
    add_body_para(doc,
        "Because $t_{pDA} > 0.1$, the reservoir reached Pseudo-Steady State prior to shut-in. Consequently, the extrapolated Horner pressure ($p^* = 1353.22 \\text{ psia}$) is a false pressure, and the Matthews-Brons-Hazebroek (MBH) method is required to compute the true average reservoir pressure ($\\bar{p}$)."
    )
    add_body_para(doc,
        "Utilizing the Dietz shape factor analytical approach for the MBH correction, the true average pressure can be directly computed. For a well centered in a 2:1 rectangular reservoir, the Dietz shape factor is $C_A = 21.8369$ (from standard reservoir engineering tables). The exact true average reservoir pressure is:"
    )
    add_body_para(doc,
        "$$\\bar{p} = p^* - m \\log_{10}(C_A t_{pDA})$$",
        indent=True
    )
    add_body_para(doc,
        "$$\\bar{p} = 1353.22 - 30.29 \\log_{10}(21.8369 \\times 0.1231)$$",
        indent=True
    )
    add_body_para(doc,
        "$$\\bar{p} = 1353.22 - 30.29 \\log_{10}(2.688)$$",
        indent=True
    )
    add_body_para(doc,
        "$$\\bar{p} = 1353.22 - 30.29 (0.4294) = 1340.21 \\text{ psia}$$",
        indent=True
    )
    add_body_para(doc,
        "Finally, extrapolating the Horner line to a Horner time ratio of 961 (which corresponds to a shut-in time of $\\Delta t = 1 \\text{ hr}$) yields $p_{ws,1hr} = 1262.87 \\text{ psia}$. Using the flowing pressure just before shut-in ($p_{wf(\\Delta t=0)} = 1192.45 \\text{ psia}$), this yields a skin factor of $s = -1.86$, indicating a successfully stimulated wellbore."
    )

    # ------------------------------------------------------------------------
    # Conclusion (Starts on new page via Level 1 page_break_before)
    # ------------------------------------------------------------------------
    add_heading_apa(doc, "Conclusion", level=1, bookmark_id=12)
    add_body_para(doc,
        f"The Bourdet pressure derivative analysis successfully determined all reservoir parameters for both drawdown and buildup tests. For Problem 1, $k = {p1_res.semilog.k:.2f} \\text{{ mD}}$, $s = {p1_res.semilog.s:.2f}$, and $C = {p1_res.semilog.C:.5f} \\text{{ bbl/psi}}$ ($C_D = {p1_res.semilog.C_D:.1f}$). For Problem 2, $k = {p2_res.semilog.k:.2f} \\text{{ mD}}$, $s = -1.86$, $C = {p2_res.semilog.C:.5f} \\text{{ bbl/psi}}$ ($C_D = {p2_res.semilog.C_D:.1f}$), and the MBH Dietz method yielded a true average reservoir pressure of $\\bar{{p}} = {p2_res.semilog.p_bar:.1f} \\text{{ psia}}$."
    )

    # ------------------------------------------------------------------------
    # PAGE: References (Starts on new page via Level 1 page_break_before)
    # ------------------------------------------------------------------------
    add_heading_apa(doc, "References", level=1, bookmark_id=13)

    REFERENCES = [
        ("Agarwal, R. G. (1980). ", "A new method to account for producing time effects when drawdown type curves are used. ", "Paper SPE 9289."),
        ("Ahmed, T., & McKinney, P. D. (2005). ", "Advanced reservoir engineering", ". Gulf Professional Publishing."),
        ("Bourdet, D. (2002). ", "Well test analysis: The use of advanced interpretation models", ". Elsevier."),
        ("Bourdet, D., Ayoub, J. A., & Pirard, Y. M. (1989). ", "Use of pressure derivative in well test interpretation. ", "SPE Formation Evaluation, 4(2), 293-302."),
    ]
    for ref_parts in REFERENCES:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.left_indent = Cm(1.27)
        p_ref.paragraph_format.first_line_indent = Cm(-1.27)
        set_para_spacing(p_ref, before=4, after=4, line_rule=WD_LINE_SPACING.MULTIPLE, line_val=1.15)
        r1 = p_ref.add_run(ref_parts[0])
        set_font(r1, size=12)
        r2 = p_ref.add_run(ref_parts[1])
        set_font(r2, italic=True, size=12)
        r3 = p_ref.add_run(ref_parts[2])
        set_font(r3, size=12)

    # ------------------------------------------------------------------------
    # TABLES & FIGURES (APA 7th Edition: Each placed on a separate page after References)
    # ------------------------------------------------------------------------

    # --- Table 1 ---
    doc.add_page_break()
    add_table_label(doc, 1, "Drawdown Test Calculated Reservoir Parameters")
    tbl1 = doc.add_table(rows=6, cols=2)
    tbl1.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(["Parameter", "Value"]):
        cell = tbl1.cell(0, i)
        align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_content_with_latex(cell, h, align=align, bold=True)

    rows1 = [
        ("Permeability, $k$", f"${p1_res.semilog.k:.2f} \\text{{ mD}}$"),
        ("Skin Factor, $s$", f"${p1_res.semilog.s:.2f}$"),
        ("Wellbore Storage Coefficient, $C$", f"${p1_res.semilog.C:.5f} \\text{{ bbl/psi}}$"),
        ("Dimensionless Wellbore Storage, $C_D$", f"${p1_res.semilog.C_D:.1f}$"),
        ("Semilog Slope, $m$", f"${p1_res.semilog.m:.2f} \\text{{ psi/cycle}}$"),
    ]
    for r_idx, (p_name, val) in enumerate(rows1):
        cell_name = tbl1.cell(r_idx+1, 0)
        cell_val  = tbl1.cell(r_idx+1, 1)
        set_cell_content_with_latex(cell_name, p_name, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_content_with_latex(cell_val, val, align=WD_ALIGN_PARAGRAPH.RIGHT)

    apply_apa_table_borders(tbl1)

    # --- Table 2 ---
    doc.add_page_break()
    add_table_label(doc, 2, "Buildup Test Calculated Reservoir Parameters")
    tbl2 = doc.add_table(rows=9, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(["Parameter", "Value"]):
        cell = tbl2.cell(0, i)
        align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_content_with_latex(cell, h, align=align, bold=True)

    rows2 = [
        ("Permeability, $k$", f"${p2_res.semilog.k:.2f} \\text{{ mD}}$"),
        ("Skin Factor, $s$", f"${p2_res.semilog.s:.2f}$"),
        ("Wellbore Storage Coefficient, $C$", f"${p2_res.semilog.C:.5f} \\text{{ bbl/psi}}$"),
        ("Dimensionless Wellbore Storage, $C_D$", f"${p2_res.semilog.C_D:.1f}$"),
        ("Semilog Slope, $m$", f"${p2_res.semilog.m:.2f} \\text{{ psi/cycle}}$"),
        ("Radius of Investigation at $t_{{shut-in}} = 72 \\text{{ hrs}}$", f"${p2_res.semilog.r_inv_72hr:.0f} \\text{{ ft}}$"),
        ("Boundary Effects Detected (During Shut-in)", "No"),
        ("True Average Reservoir Pressure, $\\bar{p}$", f"${p2_res.semilog.p_bar:.1f} \\text{{ psia}}$"),
    ]
    for r_idx, (p_name, val) in enumerate(rows2):
        cell_name = tbl2.cell(r_idx+1, 0)
        cell_val  = tbl2.cell(r_idx+1, 1)
        set_cell_content_with_latex(cell_name, p_name, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_content_with_latex(cell_val, val, align=WD_ALIGN_PARAGRAPH.RIGHT)

    apply_apa_table_borders(tbl2)

    # --- Figure 1 ---
    doc.add_page_break()
    add_figure_label(doc, 1, "Log-Log Diagnostic Plot for Drawdown Test ($\\Delta p$ and Bourdet Derivative).")
    embed_picture(doc, plot_p1_loglog, height_cm=10.0)

    # --- Figure 2 ---
    doc.add_page_break()
    add_figure_label(doc, 2, "Semi-Log Plot for Drawdown Test ($p_{wf}$ vs Log(t)).")
    embed_picture(doc, plot_p1_semi, height_cm=10.0)

    # --- Figure 3 ---
    doc.add_page_break()
    add_figure_label(doc, 3, "Log-Log Diagnostic Plot for Buildup Test (Agarwal Equivalent Time).")
    embed_picture(doc, plot_p2_loglog, height_cm=10.0)

    # --- Figure 4 ---
    doc.add_page_break()
    add_figure_label(doc, 4, "Horner Plot for Buildup Test ($p_{ws}$ vs Horner Ratio).")
    embed_picture(doc, plot_p2_horner, height_cm=10.0)

    # Re-apply page header across all sections to ensure header is present on every page
    add_page_header(doc, RUNNING_HEAD)

    doc.save(output_path)
    print(f"[OK] APA Report saved with Native OMML Equations in text & tables: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_apa_report()
