import docx
from docx.oxml.ns import qn

def audit_tables():
    doc = docx.Document("PFB2073_WellTest_Report_APA.docx")
    print(f"Total tables found: {len(doc.tables)}")

    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- Checking Table {t_idx+1} ---")
        
        # Check cell shading
        shaded_cells = 0
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.tcPr
                if tcPr is not None:
                    shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                    if shd is not None:
                        fill = shd.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '')
                        val = shd.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')
                        if fill and fill.lower() not in ['auto', 'none', 'ffffff', '00000000']:
                            shaded_cells += 1
                            print(f"  Cell ({r_idx}, {c_idx}) shaded with fill='{fill}', val='{val}'")
        print(f"  Shaded cells count: {shaded_cells}")

        # Check vertical borders (table level + cell level)
        vertical_borders_found = 0
        tblPr = table._tbl.tblPr
        tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders') if tblPr is not None else None
        if tblBorders is not None:
            for side in ['left', 'right', 'insideV']:
                b = tblBorders.find(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{side}')
                if b is not None:
                    val = b.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if val not in [None, 'none', 'nil']:
                        vertical_borders_found += 1
                        print(f"  Table-level vertical border found: side={side}, val={val}")

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.tcPr
                tcBorders = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders') if tcPr is not None else None
                if tcBorders is not None:
                    for side in ['left', 'right', 'insideV']:
                        b = tcBorders.find(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{side}')
                        if b is not None:
                            val = b.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            if val not in [None, 'none', 'nil']:
                                vertical_borders_found += 1
                                print(f"  Cell ({r_idx}, {c_idx}) vertical border found: side={side}, val={val}")
        print(f"  Vertical borders count: {vertical_borders_found}")

        # Check horizontal rules
        # Top of row 0, bottom of row 0, bottom of last row
        r0_top = False
        r0_bottom = False
        last_bottom = False

        r0_c0_tcPr = table.rows[0].cells[0]._tc.tcPr
        r0_c0_tcBorders = r0_c0_tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders') if r0_c0_tcPr is not None else None
        if r0_c0_tcBorders is not None:
            top_b = r0_c0_tcBorders.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}top')
            if top_b is not None and top_b.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') not in [None, 'none', 'nil']:
                r0_top = True
            bot_b = r0_c0_tcBorders.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom')
            if bot_b is not None and bot_b.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') not in [None, 'none', 'nil']:
                r0_bottom = True

        last_r_c0_tcPr = table.rows[-1].cells[0]._tc.tcPr
        last_r_c0_tcBorders = last_r_c0_tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders') if last_r_c0_tcPr is not None else None
        if last_r_c0_tcBorders is not None:
            bot_b = last_r_c0_tcBorders.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom')
            if bot_b is not None and bot_b.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') not in [None, 'none', 'nil']:
                last_bottom = True

        print(f"  Horizontal rule top of header: {r0_top}")
        print(f"  Horizontal rule bottom of header: {r0_bottom}")
        print(f"  Horizontal rule bottom of table: {last_bottom}")

if __name__ == '__main__':
    audit_tables()
