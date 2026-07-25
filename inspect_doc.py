import docx
import openpyxl
import re
import os

doc_path = 'PFB2073_WellTest_Report_APA.docx'
doc = docx.Document(doc_path)

print('=== SECTIONS & HEADERS / FOOTERS ===')
for i, section in enumerate(doc.sections):
    print(f'Section {i}:')
    header = section.header
    print('  Header is linked to previous:', section.header.is_linked_to_previous)
    for p in header.paragraphs:
        print(f'  Header Paragraph text: "{p.text}"')
        for r in p.runs:
            print(f'    Run text: "{r.text}"')
        # Check XML elements in header paragraph (page numbers, fields, etc)
        xml = p._p.xml
        if 'PAGE' in xml or 'NUMPAGES' in xml:
            print('    [Contains PAGE/NUMPAGES field code]')
    
    footer = section.footer
    print('  Footer is linked to previous:', section.footer.is_linked_to_previous)
    for p in footer.paragraphs:
        print(f'  Footer Paragraph text: "{p.text}"')

print('\n=== DOCUMENT PARAGRAPHS OVERVIEW ===')
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    # Print headings, table labels, equations, etc.
    if p.style.name.startswith('Heading') or 'Table' in text or '$' in text or 'Problem' in text or 'Equation' in text:
        print(f'P{idx} [{p.style.name}]: {text[:100]}')

print('\n=== TABLES ANALYSIS ===')
print(f'Total tables: {len(doc.tables)}')
for idx, table in enumerate(doc.tables):
    print(f'\n--- Table {idx+1} ---')
    print('Style:', table.style.name if table.style else 'No style')
    
    # Print preceding paragraph(s) to check table label and title above table
    # Finding paragraph before table
    tbl_elem = table._tbl
    prev_elem = tbl_elem.getprevious()
    prev_texts = []
    while prev_elem is not None:
        if prev_elem.tag.endswith('p'):
            p_obj = docx.text.paragraph.Paragraph(prev_elem, doc)
            prev_texts.append((p_obj.style.name, p_obj.text, [r.italic for r in p_obj.runs], [r.bold for r in p_obj.runs]))
            if len(prev_texts) >= 3:
                break
        prev_elem = prev_elem.getprevious()
    print('Preceding paragraphs (reverse order):')
    for pt in prev_texts:
        print('  ', pt)

    # Check table rows, headers, bold status, cell borders
    for r_idx, row in enumerate(table.rows):
        cells_text = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
        bolds = []
        for cell in row.cells:
            cell_bold = any(run.bold for p in cell.paragraphs for run in p.runs)
            bolds.append(cell_bold)
        print(f'  Row {r_idx} (bold={bolds}): {cells_text}')

    # Check table borders via XML
    tblPr = table._tbl.tblPr
    print('  tblPr borders XML:', tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders') is not None)
    tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if tblBorders is not None:
        for child in tblBorders:
            tag = child.tag.split('}')[-1]
            val = child.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', child.attrib.get('val'))
            sz = child.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', child.attrib.get('sz'))
            print(f'    Border {tag}: val={val}, sz={sz}')

print('\n=== EQUATIONS & LATEX UNITS ANALYSIS ===')
# Check all text for equations and physical units formatting
units_pattern = re.compile(r'(\b\d+(\.\d+)?\s*(md|psi|psia|bbl/psi|stb/d|STB/D|ft|hr|cp|c_t|psi\^-1)\b)')
dollar_units_pattern = re.compile(r'\$[^$]*\b(md|psi|psia|bbl/psi|stb/d|STB/D|ft|hr|cp)\b[^$]*\$')

for idx, p in enumerate(doc.paragraphs):
    text = p.text
    if any(u in text for u in ['md', 'psi', 'psia', 'bbl/psi', 'stb/d', 'STB/D', 'ft', 'hr', 'cp']):
        # Find lines that contain equations or unit references
        if '=' in text or 'equation' in text.lower() or '$' in text:
            print(f'P{idx}: {text}')

