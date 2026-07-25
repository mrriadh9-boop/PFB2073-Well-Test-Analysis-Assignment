import docx
import re

doc = docx.Document("PFB2073_WellTest_Report_APA.docx")

print("=== DEEP AUDIT: TASK 4 - PHYSICAL UNITS IN EQUATIONS & TEXT ===")

# We want to check:
# 1. Equations containing physical units: Are physical units inside LaTeX dollar signs with \text{...}?
# 2. Text containing parameters or values with physical units: Are they enclosed in $...$?

units_kw = ['STB/day', 'STB/D', 'stb/d', 'mD', 'md', 'psi', 'psia', 'bbl/psi', 'ft', 'hr', 'hrs', 'cP', 'cp', 'psi^-1']

dollar_maths = []
all_equations = []

for p_idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue

    # Find block equations ($$...$$) or inline math ($...$)
    blocks = re.findall(r'\$\$(.*?)\$\$', text)
    inlines = re.findall(r'(?<!\$)\$([^\$]+)\$(?!\$)', text)
    
    if blocks:
        for b in blocks:
            all_equations.append((p_idx, 'block', b))
    if inlines:
        for inl in inlines:
            all_equations.append((p_idx, 'inline', inl))

print(f"Total LaTeX math blocks/inlines extracted: {len(all_equations)}")

print("\n--- Equations Containing Units ---")
eq_with_units = 0
for p_idx, eq_type, eq in all_equations:
    has_unit = any(u in eq for u in units_kw)
    if has_unit:
        eq_with_units += 1
        print(f"P{p_idx} [{eq_type}]: ${eq}$")
        # Check if \text{...} or \mathrm{...} is used for units in LaTeX math
        has_text_macro = r'\text{' in eq or r'\mathrm{' in eq
        if not has_text_macro:
            print(f"  --> WARNING: Unit in LaTeX math does NOT use \\text{{...}} or \\mathrm{{...}}: {eq}")

print(f"\nTotal equations containing units: {eq_with_units}")

print("\n--- Checking Non-Math Text for Unenclosed Units in Parameter Definitions / Equations ---")
# Check if any paragraph has physical numbers/units where units are NOT inside $...$
for p_idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    # Replace all $...$ with placeholder to see what remains outside
    cleaned_text = re.sub(r'\$\$[^\$]+\$\$', '[MATHBLOCK]', text)
    cleaned_text = re.sub(r'\$[^\$]+\$', '[MATHINLINE]', cleaned_text)
    
    # Search for leftover numbers + units in cleaned_text
    matches = re.findall(r'\b\d+(\.\d+)?\s*(STB/day|STB/D|stb/d|mD|md|psia|psi|bbl/psi|ft|hrs|hr|cP|cp)\b', cleaned_text)
    if matches:
        print(f"P{p_idx} HAS UNENCLOSED UNITS IN TEXT:")
        print(f"  Original: {text}")
        print(f"  Cleaned:  {cleaned_text}")
        print(f"  Unenclosed matches: {matches}")

