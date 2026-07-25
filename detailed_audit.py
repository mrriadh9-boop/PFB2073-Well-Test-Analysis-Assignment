import os
import re
import docx
import openpyxl
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

print("==================================================")
print("AUDIT REPORT - R1-R4 DELIVERABLE COMPLIANCE")
print("==================================================\n")

# ---------------------------------------------------------
# 1. FILE EXISTENCE CHECK
# ---------------------------------------------------------
files_to_check = [
    "Problem1_LogLog_Diagnostic.png",
    "Problem1_SemiLog.png",
    "Problem2_LogLog_Diagnostic.png",
    "Problem2_Horner.png",
    "PFB2073_WellTest_Results.xlsx",
    "PFB2073_WellTest_Report_APA.docx"
]

print("--- TASK 1: Deliverable Existence Check ---")
all_exist = True
for f in files_to_check:
    exists = os.path.exists(f)
    size = os.path.getsize(f) if exists else 0
    print(f"File '{f}': {'EXISTS' if exists else 'MISSING'} ({size} bytes)")
    if not exists:
        all_exist = False

# ---------------------------------------------------------
# 2. PLOT LEGEND COMPLIANCE & OVERLAP CHECK
# ---------------------------------------------------------
print("\n--- TASK 2: Plot Legends Check ---")
# Inspect src/plotting.py to see exact legend settings
with open("src/plotting.py", "r", encoding="utf-8") as f:
    plotting_code = f.read()

import matplotlib.image as mpimg

plots = {
    "Problem1_LogLog_Diagnostic.png": "plot_drawdown_diagnostic",
    "Problem1_SemiLog.png": "plot_drawdown_semilog",
    "Problem2_LogLog_Diagnostic.png": "plot_buildup_diagnostic",
    "Problem2_Horner.png": "plot_buildup_horner",
}

for plot_file, func_name in plots.items():
    print(f"\nAnalyzing '{plot_file}' ({func_name}):")
    # Search for legend loc in plotting_code for this function
    func_match = re.search(r'def ' + func_name + r'.*?(?=def |\Z)', plotting_code, re.DOTALL)
    if func_match:
        func_body = func_match.group(0)
        loc_match = re.search(r"ax\.legend\((.*?)\)", func_body)
        if loc_match:
            legend_args = loc_match.group(1)
            print(f"  Legend call in code: ax.legend({legend_args})")
            if "loc='upper left'" in legend_args:
                print("  Legend Position Code: upper left")
            elif "loc='upper right'" in legend_args:
                print("  Legend Position Code: UPPER RIGHT (VIOLATION: Expected upper left!)")
            else:
                print(f"  Legend Position Code: {legend_args}")
        else:
            print("  Legend call not found in function body!")
    
    # Check visual overlap / bounds (using image inspection)
    img = mpimg.imread(plot_file)
    print(f"  Image shape: {img.shape} (Height x Width x Channels)")

# ---------------------------------------------------------
# 3. EXCEL WORKBOOK INSPECTION
# ---------------------------------------------------------
print("\n--- TASK 1 (Excel): PFB2073_WellTest_Results.xlsx Inspection ---")
wb = openpyxl.load_workbook("PFB2073_WellTest_Results.xlsx", data_only=True)
print(f"Workbook Sheet Names: {wb.sheetnames}")
for sheet_name in wb.sheetnames:
    ws = wb[sheet_name]
    print(f"  Sheet '{sheet_name}': max_row={ws.max_row}, max_column={ws.max_column}")
    # Print sample top-left 5x5 cell values
    sample = []
    for r in range(1, min(6, ws.max_row + 1)):
        row_vals = [ws.cell(row=r, column=c).value for c in range(1, min(6, ws.max_column + 1))]
        sample.append(row_vals)
    print(f"  Sample top-left cell values: {sample[:2]}")

# ---------------------------------------------------------
# 4. WORD DOCUMENT TABLES (APA 7th EDITION GUIDELINES)
# ---------------------------------------------------------
print("\n--- TASK 3: Word Document Tables (APA 7th) Check ---")
doc = docx.Document("PFB2073_WellTest_Report_APA.docx")

print(f"Total tables found: {len(doc.tables)}")

for t_idx, table in enumerate(doc.tables):
    print(f"\nEvaluating Table {t_idx+1}:")
    
    # Check preceding paragraphs for Table label & title
    tbl_elem = table._tbl
    prev_elem = tbl_elem.getprevious()
    preceding_paragraphs = []
    while prev_elem is not None:
        if prev_elem.tag.endswith('p'):
            p_obj = docx.text.paragraph.Paragraph(prev_elem, doc)
            preceding_paragraphs.append(p_obj)
            if len(preceding_paragraphs) >= 2:
                break
        prev_elem = prev_elem.getprevious()
    
    # Needs at least 2 preceding paragraphs (Label: "Table X", Title: "*Title*")
    print(f"  Preceding Paragraphs Count: {len(preceding_paragraphs)}")
    if len(preceding_paragraphs) >= 2:
        title_p = preceding_paragraphs[0] # paragraph immediately above table
        label_p = preceding_paragraphs[1] # paragraph above title
        
        print(f"  Label paragraph text: '{label_p.text}'")
        label_bold = any(r.bold for r in label_p.runs)
        label_italic = any(r.italic for r in label_p.runs)
        print(f"    Label Bold: {label_bold}, Italic: {label_italic}")
        
        print(f"  Title paragraph text: '{title_p.text}'")
        title_bold = any(r.bold for r in title_p.runs)
        title_italic = any(r.italic for r in title_p.runs)
        print(f"    Title Bold: {title_bold}, Italic: {title_italic}")
    elif len(preceding_paragraphs) == 1:
        p0 = preceding_paragraphs[0]
        print(f"  Single preceding paragraph text: '{p0.text}'")

    # Check Headers (Row 0 bolding)
    header_row = table.rows[0]
    header_bolds = [all(r.bold for r in cell.paragraphs[0].runs) if cell.paragraphs[0].runs else False for cell in header_row.cells]
    print(f"  Header Row Cell Bolds: {header_bolds}")
    
    # Check Borders (XML analysis of borders)
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    
    # Also check cell-level borders if tblBorders not on tblPr
    cell_borders_found = False
    top_b, bottom_b, insideH_b, left_b, right_b, insideV_b = None, None, None, None, None, None
    if tblBorders is not None:
        top_b = tblBorders.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}top')
        bottom_b = tblBorders.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom')
        insideH_b = tblBorders.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}insideH')
        left_b = tblBorders.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left')
        right_b = tblBorders.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}right')
        insideV_b = tblBorders.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}insideV')
    
    print("  Table Borders XML Check:")
    print(f"    Top border: {top_b.attrib if top_b is not None else 'None'}")
    print(f"    Bottom border: {bottom_b.attrib if bottom_b is not None else 'None'}")
    print(f"    Inside Horizontal border: {insideH_b.attrib if insideH_b is not None else 'None'}")
    print(f"    Left border: {left_b.attrib if left_b is not None else 'None'}")
    print(f"    Right border: {right_b.attrib if right_b is not None else 'None'}")
    print(f"    Inside Vertical border: {insideV_b.attrib if insideV_b is not None else 'None'}")

    # Inspect cell-level borders across all cells
    has_vertical_cell_border = False
    for r in table.rows:
        for c in r.cells:
            tcPr = c._tc.tcPr
            tcBorders = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders') if tcPr is not None else None
            if tcBorders is not None:
                for side in ['left', 'right', 'insideV']:
                    b = tcBorders.find(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{side}')
                    if b is not None and b.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') != 'none':
                        has_vertical_cell_border = True
    print(f"  Has cell-level vertical borders: {has_vertical_cell_border}")

# ---------------------------------------------------------
# 5. PHYSICAL UNITS IN LATEX DOLLAR SIGNS CHECK
# ---------------------------------------------------------
print("\n--- TASK 4: Physical Units in LaTeX Dollar Signs Check ---")
# Check all paragraphs in doc for inline/display math and units
units_list = ['md', 'mD', 'psi', 'psia', 'bbl/psi', 'stb/d', 'STB/day', 'STB/D', 'ft', 'hr', 'hrs', 'cp', 'cP']

paragraph_unit_issues = []
for p_idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    
    # Check if physical numbers with units exist outside dollar signs
    # Match patterns like "50 STB/day", "15.1396 mD", "3.06", "0.001314 bbl/psi", "960 hrs"
    # Find all occurrences of units in text
    for match in re.finditer(r'(\b\d+(\.\d+)?\s*(mD|md|psia|psi|bbl/psi|STB/day|STB/D|stb/d|ft|hrs|hr|cP|cp)\b)', text):
        unit_str = match.group(0)
        start_pos = match.start()
        
        # Check if start_pos is enclosed in $ ... $ or $$ ... $$
        # Count dollars before start_pos
        dollars_before = text[:start_pos].count('$')
        if dollars_before % 2 == 0: # Outside dollar signs!
            paragraph_unit_issues.append((p_idx, text, unit_str))

print(f"Found {len(paragraph_unit_issues)} physical numbers with units OUTSIDE LaTeX dollar signs.")
for p_idx, p_text, u_str in paragraph_unit_issues[:15]:
    print(f"  P{p_idx}: '{u_str}' in paragraph: \"{p_text[:80]}...\"")

# Check equations inside dollar signs for unit formatting
dollar_matches = []
for p_idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    matches = re.findall(r'(\$\$?.*?\$\$?)', text)
    for m in matches:
        dollar_matches.append((p_idx, m))

print(f"\nTotal LaTeX math expressions found: {len(dollar_matches)}")
print("Sample LaTeX math expressions:")
for p_idx, m in dollar_matches[:10]:
    print(f"  P{p_idx}: {m}")

# ---------------------------------------------------------
# 6. PAGE HEADERS & PAGE NUMBERS CHECK
# ---------------------------------------------------------
print("\n--- TASK 5: Page Headers & Running Head & Page Numbers Check ---")
for s_idx, section in enumerate(doc.sections):
    header = section.header
    print(f"Section {s_idx} Header:")
    for p in header.paragraphs:
        print(f"  Paragraph text: '{p.text}'")
        # Inspect XML for page number field
        p_xml = p._p.xml
        has_page_field = 'PAGE' in p_xml or 'w:fldSimple w:instr="PAGE"' in p_xml or 'w:instrText' in p_xml
        print(f"  Contains PAGE field in XML: {has_page_field}")
        if 'WELL TEST ANALYSIS ASSIGNMENT 2' in p.text:
            print("  Contains Running Head 'WELL TEST ANALYSIS ASSIGNMENT 2': YES")
        else:
            print("  Contains Running Head 'WELL TEST ANALYSIS ASSIGNMENT 2': NO")

