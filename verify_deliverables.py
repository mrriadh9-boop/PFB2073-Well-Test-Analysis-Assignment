import docx
import openpyxl
import os
import xml.etree.ElementTree as ET
import zipfile

def verify_word_report(docx_path):
    print("=== VERIFYING WORD REPORT ===")
    doc = docx.Document(docx_path)
    
    # 1. Header check
    header_tables = 0
    header_text_found = False
    page_num_field_found = False
    
    for section in doc.sections:
        header = section.header
        for tbl in header.tables:
            header_tables += 1
            for row in tbl.rows:
                for cell in row.cells:
                    if "WELL TEST ANALYSIS ASSIGNMENT 2" in cell.text:
                        header_text_found = True
                    # Check for PAGE field in xml
                    xml_str = cell._element.xml
                    if "PAGE" in xml_str or "w:fldChar" in xml_str:
                        page_num_field_found = True
                        
    print(f"Header tables found: {header_tables}")
    print(f"Header running head found: {header_text_found}")
    print(f"Page number field found: {page_num_field_found}")

    # 2. APA 7th Tables check
    print(f"\nDocument contains {len(doc.tables)} body tables.")
    apa_table_violations = []
    for t_idx, table in enumerate(doc.tables, 1):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                shd = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
                if shd is not None:
                    val = shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")
                    if val and val != "auto" and val.lower() != "ffffff":
                        apa_table_violations.append(f"Table {t_idx} Row {r_idx} Cell {c_idx} has shading {val}")
                        
                tcBorders = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders")
                if tcBorders is not None:
                    for border_name in ['left', 'right', 'insideV']:
                        b = tcBorders.find(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{border_name}")
                        if b is not None:
                            val = b.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                            if val and val != "nil" and val != "none":
                                apa_table_violations.append(f"Table {t_idx} Row {r_idx} Cell {c_idx} has {border_name} border {val}")
    print(f"APA Table Violations count: {len(apa_table_violations)}")
    if apa_table_violations:
        for v in apa_table_violations[:5]:
            print(f"  Violation: {v}")

    # 3. OMML equations check
    with zipfile.ZipFile(docx_path, 'r') as z:
        document_xml = z.read('word/document.xml').decode('utf-8')
        omath_count = document_xml.count('<m:oMath>') + document_xml.count('<m:oMath ')
        omath_para_count = document_xml.count('<m:oMathPara>') + document_xml.count('<m:oMathPara ')
        print(f"\nOMML <m:oMath> elements found in document.xml: {omath_count}")
        print(f"OMML <m:oMathPara> elements found in document.xml: {omath_para_count}")


def verify_excel_workbook(xlsx_path):
    print("\n=== VERIFYING EXCEL WORKBOOK ===")
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    sheet_names = wb.sheetnames
    print(f"Sheet names: {sheet_names}")
    
    req_sheets = ["Problem 1 Drawdown", "Problem 2 Buildup"]
    missing = [s for s in req_sheets if s not in sheet_names]
    if missing:
        print(f"MISSING SHEETS: {missing}")
    else:
        print("All required sheets present.")
        
    for sheet_name in req_sheets:
        if sheet_name in sheet_names:
            ws = wb[sheet_name]
            print(f"\nSheet '{sheet_name}' dimensions: {ws.dimensions}")
            # Check KPI values
            kpis = {}
            for r in range(1, 20):
                lbl1 = ws.cell(row=r, column=4).value
                val1 = ws.cell(row=r, column=5).value
                if lbl1:
                    kpis[str(lbl1)] = val1
            print(f"KPIs extracted from sheet '{sheet_name}':")
            for k, v in list(kpis.items())[:6]:
                print(f"  {k}: {v}")

if __name__ == "__main__":
    docx_p = r"C:\Users\60163\Downloads\PFB2073_Well_Test_Analysis_-_May_2026_1784922163\PFB2073_WellTest_Report_APA.docx"
    xlsx_p = r"C:\Users\60163\Downloads\PFB2073_Well_Test_Analysis_-_May_2026_1784922163\PFB2073_WellTest_Results.xlsx"
    verify_word_report(docx_p)
    verify_excel_workbook(xlsx_p)
