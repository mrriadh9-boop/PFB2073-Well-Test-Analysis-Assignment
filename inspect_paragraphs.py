import docx

doc = docx.Document("PFB2073_WellTest_Report_APA.docx")

p13 = doc.paragraphs[12] # 0-indexed index 12 = Paragraph 13
print("=== PARAGRAPH 13 (Index 12) ===")
print("Text:", repr(p13.text))

p28 = doc.paragraphs[27] # 0-indexed index 27 = Paragraph 28
print("\n=== PARAGRAPH 28 (Index 27) ===")
print("Text:", repr(p28.text))
