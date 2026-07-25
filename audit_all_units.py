import re
import docx

def find_all_units_in_docx():
    doc = docx.Document("PFB2073_WellTest_Report_APA.docx")
    
    # Match any pattern of number (int or float or sci notation) followed by unit
    # or unit words like hr, hrs, mD, md, psi, psia, bbl/psi, STB/day, stb/d, ft, acre, acres, cP, cp
    pattern = re.compile(r'(\b\d+(\.\d+)?\s*(mD|md|psia|psi|bbl/psi|STB/day|STB/D|stb/d|RB/STB|ft|hrs|hr|cP|cp|acre|acres)\b)', re.IGNORECASE)

    print("=== CHECKING ALL PARAGRAPHS FOR UNITS OUTSIDE DOLLAR SIGNS ===")
    violations = []
    for p_idx, p in enumerate(doc.paragraphs):
        text = p.text
        if not text.strip():
            continue
        
        # Check every match
        for match in pattern.finditer(text):
            matched_unit = match.group(0)
            start = match.start()
            
            # Check if enclosed in $...$ or $$...$$
            # To be accurate, split paragraph text by '$'
            # Odd index in split means inside $, even index means outside $
            prefix = text[:start]
            dollar_count = prefix.count('$')
            if dollar_count % 2 == 0:
                violations.append((p_idx + 1, matched_unit, text))

    print(f"Total violations found in paragraphs: {len(violations)}")
    for p_num, unit_str, text in violations:
        print(f"Paragraph {p_num}: '{unit_str}' outside LaTeX $ in: \"{text}\"")

    print("\n=== CHECKING ALL TABLES FOR UNITS OUTSIDE DOLLAR SIGNS ===")
    table_violations = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    text = p.text
                    if not text.strip():
                        continue
                    for match in pattern.finditer(text):
                        matched_unit = match.group(0)
                        start = match.start()
                        prefix = text[:start]
                        dollar_count = prefix.count('$')
                        if dollar_count % 2 == 0:
                            table_violations.append((t_idx+1, r_idx+1, c_idx+1, matched_unit, text))

    print(f"Total violations found in tables: {len(table_violations)}")
    for t_num, r_num, c_num, unit_str, text in table_violations:
        print(f"Table {t_num} R{r_num}C{c_num}: '{unit_str}' outside LaTeX $ in: \"{text}\"")

if __name__ == '__main__':
    find_all_units_in_docx()
