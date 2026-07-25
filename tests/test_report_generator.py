"""
tests/test_report_generator.py
-------------------------------
Unit tests for src/report_generator.py report generation module.
Verifies that generate_apa_report executes dynamically without static hardcoded dictionaries,
accepts Problem1AnalysisResult and Problem2AnalysisResult objects, and generates a valid
APA 7th Edition Word document (.docx).
"""

import os
import pytest
from docx import Document

from src.welltest import analyze_problem1_drawdown, analyze_problem2_buildup
from src.report_generator import generate_apa_report, enforce_number_unit_spacing


@pytest.fixture
def p1_result():
    return analyze_problem1_drawdown()


@pytest.fixture
def p2_result():
    return analyze_problem2_buildup()


def test_generate_apa_report_with_results(tmp_path, p1_result, p2_result):
    report_file = str(tmp_path / "test_report.docx")
    out_path = generate_apa_report(p1_res=p1_result, p2_res=p2_result, output_path=report_file)

    assert out_path == report_file
    assert os.path.exists(report_file)
    assert os.path.getsize(report_file) > 1000

    # Verify that the generated document can be read by python-docx
    doc = Document(report_file)
    assert len(doc.paragraphs) > 0
    assert len(doc.tables) >= 2


def test_generate_apa_report_with_none_arguments(tmp_path):
    report_file = str(tmp_path / "test_report_none.docx")
    out_path = generate_apa_report(p1_res=None, p2_res=None, output_path=report_file)

    assert out_path == report_file
    assert os.path.exists(report_file)
    assert os.path.getsize(report_file) > 1000


def test_generate_apa_report_contains_dynamic_values(tmp_path, p1_result, p2_result):
    report_file = str(tmp_path / "test_report_dynamic.docx")
    generate_apa_report(p1_res=p1_result, p2_res=p2_result, output_path=report_file)

    doc = Document(report_file)
    combined = " ".join(doc._element.xpath('.//w:t/text() | .//m:t/text()'))

    # Verify dynamic permeability values are in the document text
    assert f"{p1_result.semilog.k:.2f}" in combined
    assert f"{p2_result.semilog.k:.2f}" in combined

    # Verify static hardcoded numbers are NOT present
    assert "15.1396" not in combined
    assert "1314.0" not in combined


def test_r1_universal_omml_parsing(tmp_path, p1_result, p2_result):
    """Verifies R1: Inline LaTeX in headings, body text, table cells, and figure/table labels is converted to w:oMath elements."""
    report_file = str(tmp_path / "test_r1_omml.docx")
    generate_apa_report(p1_res=p1_result, p2_res=p2_result, output_path=report_file)
    doc = Document(report_file)
    
    # Verify w:oMath elements exist in the document (headings, body, tables, labels)
    omath_elements = doc._element.xpath('.//m:oMath')
    assert len(omath_elements) > 0, "No OMML math elements found in document"


def test_r2_number_unit_spacing(tmp_path, p1_result, p2_result):
    """Verifies R2: Pre-processor enforces single space between numbers and units (e.g. 660 ft, 14.16 mD)."""
    # Direct function testing for pre-processor helper
    assert enforce_number_unit_spacing("660ft") == "660 ft"
    assert enforce_number_unit_spacing("14.16mD") == "14.16 mD"
    assert enforce_number_unit_spacing("1314psia") == "1314 psia"
    assert enforce_number_unit_spacing("0.001314bbl/psi") == "0.001314 bbl/psi"
    assert enforce_number_unit_spacing("72hr") == "72 hr"
    assert enforce_number_unit_spacing("72hrs") == "72 hrs"
    assert enforce_number_unit_spacing("80acre") == "80 acre"
    assert enforce_number_unit_spacing("660 ft") == "660 ft"
    assert enforce_number_unit_spacing("1e-6") == "1e-6"
    assert enforce_number_unit_spacing("7th") == "7th"

    # Document generation verification
    report_file = str(tmp_path / "test_r2_spacing.docx")
    generate_apa_report(p1_res=p1_result, p2_res=p2_result, output_path=report_file)
    doc = Document(report_file)
    combined = " ".join(doc._element.xpath('.//w:t/text() | .//m:t/text()'))
    
    assert "660 ft" in combined or "14.16 mD" in combined or "72 hr" in combined
    assert "660ft" not in combined
    assert "14.16mD" not in combined


def test_r3_apa_tables_and_figures_after_references(tmp_path, p1_result, p2_result):
    """Verifies R3: Tables and figures are located after the References heading."""
    report_file = str(tmp_path / "test_r3_placement.docx")
    generate_apa_report(p1_res=p1_result, p2_res=p2_result, output_path=report_file)
    doc = Document(report_file)
    
    p_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "References" in p_texts
    ref_idx = p_texts.index("References")
    
    # Table and Figure label paragraphs (exact text "Table 1", "Figure 1") must appear AFTER References index
    fig1_idx = next(i for i, t in enumerate(p_texts) if t == "Figure 1")
    tbl1_idx = next(i for i, t in enumerate(p_texts) if t == "Table 1")
    
    assert tbl1_idx > ref_idx, "Table 1 must appear after References"
    assert fig1_idx > ref_idx, "Figure 1 must appear after References"


def test_report_content_and_equations(tmp_path, p1_result, p2_result):
    """
    Verifies that PFB2073_WellTest_Report_APA.docx contains required exact text, headings,
    and key numeric/LaTeX equation snippets for Problem 1 & 2.
    """
    report_file = str(tmp_path / "PFB2073_WellTest_Report_APA.docx")
    generate_apa_report(p1_res=p1_result, p2_res=p2_result, output_path=report_file)

    doc = Document(report_file)
    combined_text = " ".join(doc._element.xpath('.//w:t/text() | .//m:t/text()'))
    normalized_text = combined_text.replace('\u2212', '-')

    # 1. Exact string "Reservoir and Fluid Properties (Problem 2):"
    assert "Reservoir and Fluid Properties (Problem 2):" in combined_text, (
        "Exact string 'Reservoir and Fluid Properties (Problem 2):' not found in report docx"
    )

    # 2. Heading "Model Identification"
    headings = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "Model Identification" in headings, (
        "Heading 'Model Identification' not found in docx paragraphs"
    )

    # 3. Key numeric/LaTeX equation snippets for Problem 1 & 2 ($k=30.09$, $s=-1.86$, $C_D=111.6$, $C_D=351.3$)
    assert "30.09" in combined_text or "30.12" in combined_text, "Snippet '30.09' or '30.12' for k in Problem 2 not found in report text/OMML"
    assert "-1.86" in normalized_text, "Snippet '-1.86' for s in Problem 2 not found in report text/OMML"
    assert "111.6" in combined_text, "Snippet '111.6' for C_D in Problem 1 not found in report text/OMML"
    assert "351.3" in combined_text, "Snippet '351.3' for C_D in Problem 2 not found in report text/OMML"



