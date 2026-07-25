import re
import docx

def audit_units_and_headers():
    doc = docx.Document("PFB2073_WellTest_Report_APA.docx")
    
    print("=== AUDITING PAGE HEADERS ===")
    for s_idx, section in enumerate(doc.sections):
        header = section.header
        print(f"Section {s_idx+1} Header:")
        for p_idx, p in enumerate(header.paragraphs):
            text = p.text
            print(f"  Paragraph {p_idx+1} text: '{text}'")
            xml = p._p.xml
            has_page_num = 'w:fldSimple w:instr="PAGE"' in xml or 'PAGE' in xml or 'w:instrText' in xml
            print(f"  Contains PAGE field in XML: {has_page_num}")
            if 'WELL TEST ANALYSIS ASSIGNMENT 2' in text:
                print("  Contains Running Head ('WELL TEST ANALYSIS ASSIGNMENT 2'): YES")
            else:
                print("  Contains Running Head ('WELL TEST ANALYSIS ASSIGNMENT 2'): NO")

    print("\n=== AUDITING PHYSICAL UNITS IN LATEX DOLLAR SIGNS ===")
    # Common petroleum engineering units
    unit_pattern = r'(\b\d+(\.\d+)?\s*(mD|md|psia|psi|bbl/psi|STB/day|STB/D|stb/d|ft|hrs|hr|cP|cp|acre|acres)\b)'
    
    paragraphs_outside_math = []
    
    for p_idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        
        # Check all matches of numbers with units
        for match in re.finditer(unit_pattern, text):
            matched_str = match.group(0)
            start_pos = match.start()
            
            # Count dollar signs prior to start_pos
            dollars_before = text[:start_pos].count('$')
            if dollars_before % 2 == 0:  # Even number means outside dollar signs
                paragraphs_outside_math.append((p_idx + 1, matched_str, text))

    # Also check tables for units outside math
    table_units_outside_math = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if not text:
                        continue
                    for match in re.finditer(unit_pattern, text):
                        matched_str = match.group(0)
                        start_pos = match.start()
                        dollars_before = text[:start_pos].count('$')
                        if dollars_before % 2 == 0:
                            table_units_outside_math.append((t_idx+1, r_idx+1, c_idx+1, matched_str, text))

    print(f"Paragraph units outside math count: {len(paragraphs_outside_math)}")
    for p_num, u_str, full_t in paragraphs_outside_math:
        print(f"  [PARAGRAPH {p_num}] Unit outside math: '{u_str}' -> Context: \"{full_t[:100]}\"")

    print(f"Table cell units outside math count: {len(table_units_outside_math)}")
    for t_num, r_num, c_num, u_str, full_t in table_units_outside_math:
        print(f"  [TABLE {t_num} R{r_num}C{c_num}] Unit outside math: '{u_str}' -> Context: \"{full_t}\"")

if __name__ == '__main__':
    audit_units_and_headers()
